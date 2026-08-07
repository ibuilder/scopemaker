"""Word (DOCX) export.

Contracts get redlined in Word, so this has to produce a real editable document
rather than a PDF with a different extension.  Inline markup is converted to
character runs (bold/italic/underline survive), the computed outline labels are
written as literal text so the numbering matches the PDF exactly, and the
footer carries a live ``PAGE of NUMPAGES`` field.
"""

from __future__ import annotations

import io
from decimal import Decimal
from html.parser import HTMLParser
from typing import Any

from docx import Document as DocxDocument
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from ...models import Scope
from ...models.scope import SectionKind
from .html import Document, DocumentLine, build_document

BODY_FONT = "Times New Roman"
BODY_SIZE = Pt(11)
LABEL_INDENT = Inches(0.35)  # additional indent per outline level


class _RunBuilder(HTMLParser):
    """Convert a fragment of inline HTML into python-docx runs."""

    def __init__(self, paragraph: Any, *, base_bold: bool = False):
        super().__init__(convert_charrefs=True)
        self.paragraph = paragraph
        self.bold = base_bold
        self.italic = False
        self.underline = False
        self._stack: list[str] = []

    def handle_starttag(self, tag: str, attrs: list) -> None:
        if tag in ("strong", "b"):
            self.bold = True
        elif tag in ("em", "i"):
            self.italic = True
        elif tag == "u":
            self.underline = True
        elif tag == "br":
            self.paragraph.add_run().add_break()
            return
        self._stack.append(tag)

    def handle_endtag(self, tag: str) -> None:
        if tag in self._stack:
            # Pop back to the matching tag so malformed nesting cannot leave
            # the whole remainder of the document bold.
            while self._stack:
                popped = self._stack.pop()
                if popped == tag:
                    break
        if tag in ("strong", "b"):
            self.bold = False
        elif tag in ("em", "i"):
            self.italic = False
        elif tag == "u":
            self.underline = False

    def handle_data(self, data: str) -> None:
        if not data:
            return
        run = self.paragraph.add_run(data)
        run.bold = self.bold
        run.italic = self.italic
        run.underline = self.underline
        run.font.name = BODY_FONT
        run.font.size = BODY_SIZE


def _add_html(paragraph: Any, html: str, *, base_bold: bool = False) -> None:
    if not html:
        return
    builder = _RunBuilder(paragraph, base_bold=base_bold)
    builder.feed(html)
    builder.close()


def _add_field(paragraph: Any, instruction: str) -> None:
    """Insert a Word field code (used for PAGE and NUMPAGES)."""
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {instruction} "

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")

    run = paragraph.add_run()
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _money(value: Decimal | None, currency: str) -> str:
    symbol = "$" if currency in ("USD", "CAD", "AUD") else ""
    if value is None:
        return f"{symbol} TBD"
    return f"{symbol}{Decimal(value):,.2f}"


def _configure_page(document: Any) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.NEW_PAGE
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    style = document.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = BODY_SIZE
    # East-Asian font mapping, otherwise Word substitutes for non-Latin glyphs.
    style.element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.0


def _add_header_footer(document: Any, doc: Document) -> None:
    header = document.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_text = " | ".join(
        part
        for part in (
            doc.project.get("name"),
            doc.bid_package.get("number"),
            doc.trade_name,
        )
        if part
    )
    run = paragraph.add_run(header_text)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix = paragraph.add_run(f"{doc.title}    Page ")
    prefix.font.size = Pt(8)
    prefix.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _add_field(paragraph, "PAGE")
    middle = paragraph.add_run(" of ")
    middle.font.size = Pt(8)
    middle.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    _add_field(paragraph, "NUMPAGES")
    for run in paragraph.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _add_title(document: Any, doc: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(doc.title.upper())
    run.bold = True
    run.underline = True
    run.font.size = Pt(13)
    run.font.name = BODY_FONT

    subtitle_parts = [p for p in (doc.division_label, doc.trade_name) if p]
    if subtitle_parts:
        sub = document.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub.add_run(" – ".join(subtitle_parts))
        run.bold = True
        run.font.size = Pt(11)


def _add_project_block(document: Any, doc: Document) -> None:
    rows = [
        ("Project", " ".join(p for p in (doc.project.get("number"),
                                         doc.project.get("name")) if p)),
        ("Location", doc.project.get("location")),
        ("Bid Package", " ".join(p for p in (doc.bid_package.get("number"),
                                             doc.bid_package.get("name")) if p)),
        ("Subcontractor", doc.bid_package.get("subcontractor")),
    ]
    rows = [(label, value) for label, value in rows if value]
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for index, (label, value) in enumerate(rows):
        left = table.cell(index, 0)
        left.width = Inches(1.4)
        run = left.paragraphs[0].add_run(f"{label}:")
        run.bold = True
        run.font.size = Pt(10)

        right = table.cell(index, 1)
        right.width = Inches(5.1)
        run = right.paragraphs[0].add_run(str(value))
        run.font.size = Pt(10)

    document.add_paragraph()


def _add_lines(document: Any, lines: list[DocumentLine], depth: int = 0) -> None:
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = LABEL_INDENT * (depth + 1)
        paragraph.paragraph_format.first_line_indent = -LABEL_INDENT
        paragraph.paragraph_format.space_after = Pt(4)

        label = paragraph.add_run(f"{line.label}\t")
        label.bold = False
        label.font.name = BODY_FONT
        label.font.size = BODY_SIZE

        _add_html(paragraph, line.html)
        _add_lines(document, line.children, depth + 1)


def _add_recap(document: Any, doc: Document) -> None:
    if not doc.recap_rows:
        return
    table = document.add_table(rows=len(doc.recap_rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, row in enumerate(doc.recap_rows):
        left = table.cell(index, 0)
        left.width = Inches(4.5)
        run = left.paragraphs[0].add_run(row.label)
        run.bold = row.is_total
        run.font.size = Pt(11)

        right = table.cell(index, 1)
        right.width = Inches(2.0)
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = right.paragraphs[0].add_run(_money(row.amount, doc.currency))
        run.bold = row.is_total
        run.font.size = Pt(11)


def render_docx(scope: Scope, *, organization: Any = None) -> bytes:
    """Render a scope to a Word document."""
    doc = build_document(scope, organization=organization)

    document = DocxDocument()
    _configure_page(document)
    _add_header_footer(document, doc)
    _add_title(document, doc)
    _add_project_block(document, doc)

    for section in doc.enabled_sections:
        heading = document.add_paragraph()
        heading.paragraph_format.space_before = Pt(10)
        heading.paragraph_format.space_after = Pt(4)
        run = heading.add_run(f"{section.number} {section.heading.upper()}")
        run.bold = True
        run.underline = True
        run.font.size = Pt(11)
        run.font.name = BODY_FONT

        if section.body_html:
            body = document.add_paragraph()
            body.paragraph_format.left_indent = LABEL_INDENT
            _add_html(body, section.body_html)

        if section.kind == SectionKind.RECAP:
            _add_recap(document, doc)
        else:
            _add_lines(document, section.lines)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
