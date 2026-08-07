"""Export formats.

The central guarantee is that every format says the same thing with the same
clause numbers, so most of these tests compare one renderer against another.
"""

from __future__ import annotations

import io
import json

import pytest

from scopemaker.services.renderers import (
    PDF_AVAILABLE,
    build_document,
    render_docx,
    render_html,
    render_json,
    render_markdown,
    render_pdf,
)
from scopemaker.services.renderers.json_export import build_payload


def flatten(lines):
    for line in lines:
        yield line
        yield from flatten(line.children)


def test_document_numbers_sections_and_items(app, scope, organization):
    doc = build_document(scope, organization=organization)
    numbers = [section.number for section in doc.enabled_sections]
    assert numbers == [f"{i}." for i in range(1, len(numbers) + 1)]

    summary = next(s for s in doc.enabled_sections if s.key == "summary")
    labels = [line.label for line in summary.lines]
    assert labels[:3] == ["1.", "2.", "3."]

    nested = [line for line in flatten(summary.lines) if line.depth == 1]
    assert nested, "spec sections did not nest"
    assert nested[0].label.startswith("3.")


def test_markdown_carries_the_same_labels(app, scope, organization):
    doc = build_document(scope, organization=organization)
    markdown = render_markdown(scope, organization=organization).decode()

    summary = next(s for s in doc.enabled_sections if s.key == "summary")
    for line in list(flatten(summary.lines))[:6]:
        assert f"{line.label} {line.text}" in markdown


def test_markdown_decodes_html_entities(app, scope, organization):
    markdown = render_markdown(scope, organization=organization).decode()
    # Spec section rows are written with an &ndash; separator.
    assert "&ndash;" not in markdown
    assert "&amp;" not in markdown
    assert "–" in markdown


def test_json_export_shape(app, scope, organization):
    payload = json.loads(render_json(scope, organization=organization).decode())
    assert payload["format"] == "scopemaker.scope"
    assert payload["scope"]["division"] == "Division 21"
    assert payload["project"]["name"] == "Riverside Medical Center"

    sections = {s["key"]: s for s in payload["sections"]}
    assert "inclusions" in sections
    summary = sections["summary"]
    assert summary["items"][0]["number"] == "1."
    assert summary["items"][0]["path"] == "1"

    nested = [i for i in summary["items"] if i.get("children")]
    assert nested, "nested spec sections missing from JSON"
    assert nested[0]["children"][0]["path"].count(".") == 1


def test_json_recap_totals(app, scope, organization):
    payload = build_payload(scope, organization=organization)
    total = next(row for row in payload["recap"] if row["is_total"])
    assert total["amount"] == "1425000.00"


def test_html_is_standalone_and_paged(app, scope, organization):
    html = render_html(scope, organization=organization, standalone=True)
    assert html.lstrip().startswith("<!DOCTYPE html>")
    # The paged-media rules must be inlined, not linked, or WeasyPrint and the
    # browser print view would resolve the stylesheet differently.
    assert "@page" in html
    assert "counter(page)" in html and "counter(pages)" in html
    assert "<link" not in html
    assert "EXHIBIT B" in html


def test_inlined_stylesheet_is_not_html_escaped(app, scope, organization):
    """Jinja autoescaping silently corrupts an inlined stylesheet.

    `"` becomes `&#34;` and `>` becomes `&gt;`, which breaks
    `content: "Page " counter(page)` in the @page margin boxes and every child
    combinator. The PDF then renders with no page numbers -- and nothing else
    complains, because CSS parse errors are only warnings.
    """
    html = render_html(scope, organization=organization, standalone=True)
    style = html[html.index("<style>") : html.index("</style>")]

    for entity in ("&#34;", "&quot;", "&gt;", "&lt;", "&amp;"):
        assert entity not in style, f"{entity} in the inlined stylesheet"

    # The rules that the escaping used to break.
    assert 'content: "Page " counter(page) " of " counter(pages);' in style
    assert "> .doc__item-label" in style


def test_running_elements_are_declared_for_the_page_margins(app, scope, organization):
    html = render_html(scope, organization=organization, standalone=True)
    assert "position: running(doc-header)" in html
    assert "position: running(doc-footer)" in html
    assert "content: element(doc-header)" in html
    assert "content: element(doc-footer)" in html


def test_html_body_fragment_is_not_a_full_page(app, scope, organization):
    body = render_html(scope, organization=organization, standalone=False)
    assert "<!DOCTYPE" not in body
    assert 'class="doc' in body


def test_docx_is_a_real_word_file(app, scope, organization):
    payload = render_docx(scope, organization=organization)
    assert payload[:2] == b"PK", "not a zip container"

    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(payload))
    text = "\n".join(p.text for p in document.paragraphs)
    assert "EXHIBIT B" in text.upper()

    doc = build_document(scope, organization=organization)
    summary = next(s for s in doc.enabled_sections if s.key == "summary")
    # Labels are written literally so Word matches the PDF exactly.
    assert any(line.label in text for line in summary.lines)


def test_docx_footer_has_a_page_number_field(app, scope, organization):
    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(render_docx(scope, organization=organization)))
    footer_xml = document.sections[0].footer.paragraphs[0]._p.xml
    assert "PAGE" in footer_xml
    assert "NUMPAGES" in footer_xml


def test_locked_scope_still_exports(app, db, scope, organization, user):
    from scopemaker.services.scope_builder import issue_scope

    issue_scope(scope, user_id=user.id)
    assert render_markdown(scope, organization=organization)
    assert render_docx(scope, organization=organization)


@pytest.mark.pdf
@pytest.mark.skipif(not PDF_AVAILABLE, reason="WeasyPrint native libraries unavailable")
def test_pdf_is_paginated_with_selectable_text(app, scope, organization):
    """The failure this guards against is the prototype's screenshot PDF.

    That produced a single raster page with no extractable text; a real
    exhibit must paginate and must be searchable.
    """
    from pypdf import PdfReader

    payload = render_pdf(scope, organization=organization)
    assert payload[:5] == b"%PDF-"

    reader = PdfReader(io.BytesIO(payload))
    assert len(reader.pages) >= 2, "a full scope should not fit on one page"

    # Justified text makes the extractor emit runs of spaces at the positions
    # where glyphs were spread apart, so compare on normalised whitespace.
    text = " ".join(
        " ".join(page.extract_text() or "" for page in reader.pages).split()
    )
    assert "EXHIBIT B" in text.upper()
    assert "SCOPE OF WORK" in text.upper()

    doc = build_document(scope, organization=organization)
    inclusions = next(s for s in doc.enabled_sections if s.key == "inclusions")
    first = inclusions.lines[0]
    # A distinctive fragment of real clause text must be extractable -- this is
    # what the prototype's rasterised PDF could never satisfy.
    fragment = " ".join(first.text.split(",")[0].split())[:40]
    assert fragment and fragment in text

    # Running footer with "Page N of M", and the running header on later pages.
    assert "Page 1 of" in text
    assert f"of {len(reader.pages)}" in text
    assert doc.project["name"] in text


@pytest.mark.pdf
@pytest.mark.skipif(not PDF_AVAILABLE, reason="WeasyPrint native libraries unavailable")
def test_pdf_page_count_grows_with_content(app, db, scope, organization):
    from scopemaker.services.renderers.pdf import render_pdf_pages

    before = render_pdf_pages(scope, organization=organization)

    from scopemaker.models import ScopeItem

    section = scope.section("inclusions")
    for index in range(60):
        db.session.add(
            ScopeItem(
                section_id=section.id,
                text_html=f"Additional obligation number {index} " + ("padding text " * 20),
                position=1000 + index,
            )
        )
    db.session.commit()

    after = render_pdf_pages(scope, organization=organization)
    assert after > before, "content overflowed instead of paginating"


def test_pdf_reports_a_clear_error_when_unavailable(app, scope, organization):
    """Without the native stack the failure must be actionable, not a crash."""
    if PDF_AVAILABLE:
        pytest.skip("native stack present")

    from scopemaker.errors import RenderError

    with pytest.raises(RenderError) as excinfo:
        render_pdf(scope, organization=organization)
    assert excinfo.value.code == "pdf_unavailable"
    assert "hint" in excinfo.value.details
